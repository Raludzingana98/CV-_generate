from docx import Document
from docx.shared import Pt

def remove_empty_bullets(doc):
    
    paragraphs_to_delete = []
    
    for paragraph in doc.paragraphs:
       
        if 'List' in paragraph.style.name:
            if not paragraph.text.strip():
                paragraphs_to_delete.append(paragraph)
    
    for p in paragraphs_to_delete:
        p._element.getparent().remove(p._element)

def remove_empty_rows(doc):
    for table in doc.tables:
        rows_to_delete = []
        for row in table.rows:
            is_empty_row = True
            for cell in row.cells:
                if cell.text.strip():
                    is_empty_row = False
                    break
            
            if is_empty_row:
                rows_to_delete.append(row)
        
        for row in rows_to_delete:
            
            row._element.getparent().remove(row._element)

def remove_empty_sections(doc):
    
    target_headers = [
        "CAREER SUMMARY", 
        "PROJECTS", 
        "QUALIFICATIONS", 
        "PROFESSIONAL MEMBERSHIPS", 
        "CAREER HISTORY"
    ]

    body_elements = list(doc.element.body.iterchildren())
    
    
    header_indices = []
    for i, element in enumerate(body_elements):
        if element.tag.endswith('p'):
            text = "".join([node.text for node in element.findall('.//w:t', namespaces=element.nsmap) if node.text])
            if text.strip() in target_headers:
                header_indices.append(i)
                
    
    for i in reversed(header_indices):
        start_index = i
        next_header_index = len(body_elements)
        
        for hi in header_indices:
            if hi > start_index:
                next_header_index = min(next_header_index, hi)
                break 

        is_section_empty = True
        elements_to_remove = [body_elements[start_index]] 
        
        for j in range(start_index + 1, next_header_index):
            elem = body_elements[j]
            elements_to_remove.append(elem)
            
            if elem.tag.endswith('p'):
                text = "".join([node.text for node in elem.findall('.//w:t', namespaces=elem.nsmap) if node.text])
                if text.strip():
                    is_section_empty = False
                    elements_to_remove = [] 
                    break
            elif elem.tag.endswith('tbl'):
                text = "".join([node.text for node in elem.findall('.//w:t', namespaces=elem.nsmap) if node.text])
                if text.strip():
                     is_section_empty = False
                     elements_to_remove = []
                     break
        
        if is_section_empty:
            
            if start_index > 0:
                prev_elem = body_elements[start_index - 1]
                
                page_breaks = prev_elem.findall('.//w:br[@w:type="page"]', namespaces=prev_elem.nsmap)
                if page_breaks:
                    elements_to_remove.insert(0, prev_elem)

            
            for elem in elements_to_remove:
                if elem.getparent() is not None:
                    elem.getparent().remove(elem)

def remove_placeholders(doc):
    import re
    
    def clean_runs(runs):
        if not runs:
            return
            
        full_text = "".join(run.text for run in runs if run.text)
        
        if '{{' in full_text:
            new_full_text = re.sub(r'\{\{.*?\}\}', '', full_text)
            if full_text != new_full_text:
                for i, run in enumerate(runs):
                    if i == 0:
                        run.text = new_full_text
                    else:
                        run.text = ""

   
    for paragraph in doc.paragraphs:
        clean_runs(paragraph.runs)
        
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                 for paragraph in cell.paragraphs:
                     clean_runs(paragraph.runs)

def clean_up_document(file_path):
    try:
        doc = Document(file_path)
        
        remove_placeholders(doc)
        remove_empty_bullets(doc)
        remove_empty_rows(doc)
        remove_empty_sections(doc)
        
        doc.save(file_path)
        print(f"Cleaned up document: {file_path}")
        return True
    except Exception as e:
        print(f"Failed to clean up document {file_path}: {e}")
        return False

if __name__ == "__main__":
    
    import sys
    if len(sys.argv) > 1:
        clean_up_document(sys.argv[1])
