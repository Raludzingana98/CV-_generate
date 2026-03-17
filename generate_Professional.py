from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import json
import os
import style_utils

# ------------------------------------------------
# FIRST PAGE (COVER PAGE)
# ------------------------------------------------
def generate_first_page(doc, data=None):
    data = data or {}
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Inches(2.3)

    # Candidate Name
    name = data.get("CANDIDATE_NAME", "{{CANDIDATE_NAME}}")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(name.upper())
    run.font.name = "Arial"
    run.font.size = Pt(36)
    run.font.bold = True

    # Professional Title
    title = data.get("PROFESSIONAL_TITLE", "{{PROFESSIONAL_TITLE}}")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.name = "Arial"
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(80, 80, 80)

    # Divider
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("____________________________________________________")

    # CV Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("CURRICULUM VITAE")
    run.font.size = Pt(14)
    run.font.bold = True

    # Position Applied For
    position = data.get("POSITION", "{{POSITION}}")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Position Applied For").bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(position)
    run.font.size = Pt(18)
    run.font.bold = True

    # Date
    date_val = data.get("DATE", "{{DATE}}")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(date_val)
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(23, 54, 93)
    p.paragraph_format.space_after = Pt(30)

    # Compensation Card
    table = doc.add_table(rows=1, cols=1)
    style_utils.apply_table_styles(table)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.width = Inches(3.5)
    cell = table.rows[0].cells[0]
    style_utils.set_cell_border(cell)
    style_utils.set_cell_shading(cell, "D9D9D9")

    def add_centered(text, bold=True):
        p = cell.add_paragraph(text)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.runs[0]
        r.font.name = "Arial"
        r.font.size = Pt(9)
        r.font.bold = bold

    add_centered("COMPENSATION OVERVIEW")
    cell.add_paragraph()
    add_centered("Previous Salary")
    add_centered(data.get("PREVIOUS_SALARY", "{{PREVIOUS_SALARY}}"), False)
    cell.add_paragraph()
    add_centered("Salary Expectation")
    add_centered(data.get("SALARY_EXPECTATION", "{{SALARY_EXPECTATION}}"), False)
    doc.add_paragraph()

# ------------------------------------------------
# SECTION RENDERING HELPERS
# ------------------------------------------------

def add_personal_details(doc, data):
    table = doc.add_table(rows=16, cols=2)
    style_utils.apply_table_styles(table)
    header_row = table.rows[0]
    header_cell = header_row.cells[0].merge(header_row.cells[1])
    header_para = header_cell.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header_para.add_run("PERSONAL DETAILS")
    run.font.name = 'Arial'
    run.font.size = Pt(9)
    run.font.bold = True
    style_utils.set_cell_shading(header_cell, 'BEBEBE')
    
    fields = [
        ("Full Name", "CANDIDATE_NAME"),
        ("Known As", "KNOWN_AS"),
        ("Race and Gender", "RACE_GENDER"),
        ("Nationality", "NATIONALITY"),
        ("Marital Status", "MARITAL_STATUS"),
        ("Date of birth", "DOB"),
        ("Driver's License / Own Car", "DRIVERS"),
        ("Notice Period", "NOTICE"),
        ("Do you have clear criminal record", "CRIMINAL"),
        ("Area of residence", "RESIDENCE"),
        ("Languages", "LANGUAGES"),
        ("Salary Expectations", "SALARY_EXP"),
        ("Skills", "SKILLS"),
        ("Computer literacy", "COMPUTER"),
        ("Experience", "EXPERIENCE_SUMMARY")
    ]
    for idx, (label, key) in enumerate(fields):
        row = table.rows[idx + 1]
        c0, c1 = row.cells
        c0.text = label
        style_utils.set_cell_shading(c0, 'F1F1F1')
        p0 = c0.paragraphs[0]
        for run in p0.runs:
            run.font.name, run.font.size, run.font.bold = 'Arial', Pt(9), True
        
        c1.text = str(data.get(key, f"{{{{{key}}}}}"))
        p1 = c1.paragraphs[0]
        for run in p1.runs:
            run.font.name, run.font.size, run.font.bold = 'Arial', Pt(9), False
    
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            style_utils.set_cell_border(cell, top=(r_idx==0), bottom=True, left=(c_idx==0), right=(c_idx==len(row.cells)-1 or (r_idx==0 and c_idx==0)))
    doc.add_paragraph()

def add_qualifications(doc, data):
    doc.add_paragraph("QUALIFICATIONS", style="Heading 1")
    quals = data.get("qualifications", [])
    num_rows = max(len(quals), 4) + 1
    table = doc.add_table(rows=num_rows, cols=4)
    style_utils.apply_table_styles(table)
    headers = ["QUALIFICATION", "INSTITUTION", "YEAR", "STATUS"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        style_utils.set_cell_shading(cell, 'BEBEBE')
        style_utils.set_cell_border(cell)
        for run in cell.paragraphs[0].runs: run.font.name, run.font.size, run.font.bold = 'Arial', Pt(9), True
            
    for i in range(1, num_rows):
        row = table.rows[i]
        if i-1 < len(quals):
            q = quals[i-1]
            vals = [q.get("qualification", ""), q.get("institution", ""), q.get("date", ""), q.get("status", "")]
        else:
            vals = [f"{{{{QUAL_{i}}}}}", f"{{{{INSTITUTION_{i}}}}}", f"{{{{YEAR_{i}}}}}", f"{{{{STATUS_{i}}}}}" ]
        for j, val in enumerate(vals):
            cell = row.cells[j]
            cell.text = str(val)
            style_utils.set_cell_border(cell)
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT if j < 2 else WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs: run.font.name, run.font.size = 'Arial', Pt(9)
        style_utils.set_cell_shading(row.cells[0], 'F1F1F1')
    doc.add_paragraph()

def add_memberships(doc, data):
    doc.add_paragraph("PROFESSIONAL MEMBERSHIPS", style="Heading 1")
    members = data.get("professional_membership", [])
    num_rows = max(len(members), 2) + 1
    table = doc.add_table(rows=num_rows, cols=4)
    style_utils.apply_table_styles(table)
    headers = ["NAME", "PROFESSION", "YEAR", "PROF NO"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        style_utils.set_cell_shading(cell, 'BEBEBE')
        style_utils.set_cell_border(cell)
        for run in cell.paragraphs[0].runs: run.font.name, run.font.size, run.font.bold = 'Arial', Pt(9), True
            
    for i in range(1, num_rows):
        row = table.rows[i]
        if i-1 < len(members):
            m = members[i-1]
            vals = [m.get("board_name", ""), m.get("profession", ""), m.get("year", ""), m.get("prof_no", "")]
        else:
            vals = [f"{{{{MEMBER_NAME_{i}}}}}", f"{{{{PROFESSION_{i}}}}}", f"{{{{MEMBER_YEAR_{i}}}}}", f"{{{{PROF_NO_{i}}}}}" ]
        for j, val in enumerate(vals):
            cell = row.cells[j]
            cell.text = str(val)
            style_utils.set_cell_border(cell)
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT if j < 2 else WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs: run.font.name, run.font.size = 'Arial', Pt(9)
        style_utils.set_cell_shading(row.cells[0], 'F1F1F1')
    doc.add_paragraph()

def add_projects(doc, data):
    doc.add_paragraph("PROJECTS", style="Heading 1")
    projects = data.get("projects", [])
    if not projects:
        projects = [{"title": f"{{{{PROJECT_TITLE_{i}}}}}", "value": f"{{{{PROJECT_VALUE_{i}}}}}", "bullets": [f"{{{{PROJECT_DESCRIPTION_{i}}}}}"]} for i in range(1, 3)]
    for proj in projects:
        title = proj.get("title", proj.get("project_name", ""))
        value = proj.get("value", proj.get("project_value", ""))
        bullets = proj.get("bullets", proj.get("description", []))
        if isinstance(bullets, str): bullets = [bullets]
        
        p_title = doc.add_paragraph()
        p_title.add_run(str(title)).font.name, p_title.runs[0].font.size, p_title.runs[0].font.bold = 'Arial', Pt(9), True
        p_title.paragraph_format.left_indent = Inches(0.25)
        p_title.paragraph_format.space_after = Pt(0)
        
        p_value = doc.add_paragraph()
        p_value.add_run(str(value)).font.name, p_value.runs[0].font.size, p_value.runs[0].font.bold = 'Arial', Pt(9), True
        p_value.paragraph_format.left_indent = Inches(0.25)
        p_value.paragraph_format.space_after = Pt(0)
        
        for b in bullets:
            pb = doc.add_paragraph(str(b), style='List Bullet')
            for run in pb.runs: run.font.name, run.font.size = 'Arial', Pt(9)
        doc.add_paragraph()

def add_career_summary(doc, data):
    doc.add_paragraph("CAREER SUMMARY", style="Heading 1")
    summary = data.get("career_summary", [])
    num_rows = max(len(summary), 5) + 1
    table = doc.add_table(rows=num_rows, cols=3)
    style_utils.apply_table_styles(table)
    headers = ["COMPANY", "POSITION", "DATES"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        style_utils.set_cell_shading(cell, 'BEBEBE')
        style_utils.set_cell_border(cell)
        for run in cell.paragraphs[0].runs: run.font.name, run.font.size, run.font.bold = 'Arial', Pt(9), True
            
    for i in range(1, num_rows):
        row = table.rows[i]
        if i-1 < len(summary):
            s = summary[i-1]
            vals = [s.get("company", ""), s.get("position", ""), s.get("dates", "")]
        else:
            vals = [f"{{{{SUMMARY_COMPANY_{i}}}}}", f"{{{{SUMMARY_POSITION_{i}}}}}", f"{{{{SUMMARY_DATES_{i}}}}}" ]
        for j, val in enumerate(vals):
            cell = row.cells[j]
            cell.text = str(val)
            style_utils.set_cell_border(cell)
            for run in cell.paragraphs[0].runs: run.font.name, run.font.size = 'Arial', Pt(9)
    doc.add_paragraph()

def add_career_history(doc, data):
    doc.add_paragraph("CAREER HISTORY", style="Heading 1")
    style_utils.build_career_history_table(doc, "CURRENT EMPLOYMENT", data.get("current_employment", []), style_variant="master")
    style_utils.build_career_history_table(doc, "PREVIOUS EMPLOYMENT", data.get("previous_employment", []), style_variant="master")

def add_highlights(doc, data):
    doc.add_paragraph("CAREER HIGHLIGHTS", style="Heading 1")
    highlights = data.get("career_highlights", [])
    if not highlights: highlights = ["{{HIGHLIGHT_1}}", "{{HIGHLIGHT_2}}", "{{HIGHLIGHT_3}}"]
    for h in highlights:
        p = doc.add_paragraph(str(h), style='List Bullet')
        for run in p.runs: run.font.name, run.font.size = 'Arial', Pt(9)
    doc.add_paragraph()

def add_course_work(doc, data):
    doc.add_paragraph("COURSE WORK", style="Heading 1")
    course_work = data.get("course_work", [])
    if not course_work: course_work = ["{{COURSE_WORK_1}}", "{{COURSE_WORK_2}}"]
    for cw in course_work:
        p = doc.add_paragraph(str(cw), style='List Bullet')
        for run in p.runs: run.font.name, run.font.size = 'Arial', Pt(9)
    doc.add_paragraph()

def add_training(doc, data):
    doc.add_paragraph("TRAINING", style="Heading 1")
    training = data.get("training", [])
    if not training: training = ["{{TRAINING_1}}", "{{TRAINING_2}}"]
    for t in training:
        p = doc.add_paragraph(str(t), style='List Bullet')
        for run in p.runs: run.font.name, run.font.size = 'Arial', Pt(9)
    doc.add_paragraph()

def add_courses(doc, data):
    doc.add_paragraph("COURSES", style="Heading 1")
    courses = data.get("courses", [])
    num_rows = max(len(courses), 4) + 1
    table = doc.add_table(rows=num_rows, cols=4)
    style_utils.apply_table_styles(table)
    headers = ["COURSE NAME", "INSTITUTION", "DATE", "STATUS"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        style_utils.set_cell_shading(cell, 'BEBEBE')
        style_utils.set_cell_border(cell)
        for run in cell.paragraphs[0].runs: run.font.name, run.font.size, run.font.bold = 'Arial', Pt(9), True
            
    for i in range(1, num_rows):
        row = table.rows[i]
        if i-1 < len(courses):
            c = courses[i-1]
            vals = [c.get("course_name", ""), c.get("institution", ""), c.get("date", ""), c.get("status", "")]
        else:
            vals = [f"{{{{COURSE_{i}}}}}", f"{{{{INSTITUTION_{i}}}}}", f"{{{{DATE_{i}}}}}", f"{{{{STATUS_{i}}}}}" ]
        for j, val in enumerate(vals):
            cell = row.cells[j]
            cell.text = str(val)
            style_utils.set_cell_border(cell)
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT if j < 2 else WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs: run.font.name, run.font.size = 'Arial', Pt(9)
        style_utils.set_cell_shading(row.cells[0], 'F1F1F1')
    doc.add_paragraph()

# ------------------------------------------------
# TEMPLATE CONFIGURATION
# ------------------------------------------------

SECTION_MAP = {
    "Personal Details": add_personal_details,
    "Qualification": add_qualifications,
    "Professional Membership": add_memberships,
    "Projects": add_projects,
    "Career Summary": add_career_summary,
    "Career History": add_career_history,
    "Career Highlights": add_highlights,
    "Course Work": add_course_work,
    "Training": add_training,
    "Courses": add_courses,
}

DATA_KEY_MAP = {
    "Personal Details": "personal_details",
    "Qualification": "qualifications",
    "Professional Membership": "professional_membership",
    "Projects": "projects",
    "Career Summary": "career_summary",
    "Career History": "career_history",
    "Career Highlights": "career_highlights",
    "Course Work": "course_work",
    "Training": "training",
    "Courses": "courses",
}


def _render_key_value_section(doc, title, fields, data):
    table = doc.add_table(rows=len(fields) + 1, cols=2)
    style_utils.apply_table_styles(table)
    header = table.rows[0].cells[0].merge(table.rows[0].cells[1])
    header_para = header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header_para.add_run(title.upper())
    run.font.name, run.font.size, run.font.bold = 'Arial', Pt(9), True
    style_utils.set_cell_shading(header, 'BEBEBE')
    for idx, field in enumerate(fields):
        label = field.get("label") or field.get("name") or ""
        data_key = field.get("key") or field.get("data_key")
        val = style_utils._get_nested_value(data, data_key, f"{{{{{data_key}}}}}")
        row = table.rows[idx + 1]
        c0, c1 = row.cells
        c0.text = label
        style_utils.set_cell_shading(c0, 'F1F1F1')
        p0 = c0.paragraphs[0]
        for run in p0.runs: run.font.name, run.font.size, run.font.bold = 'Arial', Pt(9), True
        c1.text = str(val)
        p1 = c1.paragraphs[0]
        for run in p1.runs: run.font.name, run.font.size, run.font.bold = 'Arial', Pt(9), False
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            style_utils.set_cell_border(cell, top=(r_idx == 0), bottom=True, left=(c_idx == 0), right=(c_idx == len(row.cells) - 1 or (r_idx == 0 and c_idx == 0)))
    doc.add_paragraph()

def _render_table_section(doc, title, list_data, columns, min_rows=0):
    doc.add_paragraph(title.upper(), style="Heading 1")
    rows = max(len(list_data), min_rows) + 1
    table = doc.add_table(rows=rows, cols=len(columns))
    style_utils.apply_table_styles(table)
    for i, col in enumerate(columns):
        cell = table.rows[0].cells[i]
        cell.text = col.get("header", "")
        style_utils.set_cell_shading(cell, 'BEBEBE')
        style_utils.set_cell_border(cell)
        for run in cell.paragraphs[0].runs: run.font.name, run.font.size, run.font.bold = 'Arial', Pt(9), True
    for i in range(1, rows):
        row = table.rows[i]
        item = list_data[i - 1] if i - 1 < len(list_data) else {}
        for j, col in enumerate(columns):
            key = col.get("key")
            val = style_utils._get_nested_value(item, key, "")
            if val == "" and "placeholder" in col: val = col["placeholder"].format(i=i)
            cell = row.cells[j]
            cell.text = str(val)
            style_utils.set_cell_border(cell)
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT if j < 2 else WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs: run.font.name, run.font.size = 'Arial', Pt(9)
        style_utils.set_cell_shading(row.cells[0], 'F1F1F1')
    doc.add_paragraph()

def _render_list_section(doc, title, items):
    doc.add_paragraph(title.upper(), style="Heading 1")
    for v in (items or []):
        p = doc.add_paragraph(str(v), style='List Bullet')
        for run in p.runs: run.font.name, run.font.size = 'Arial', Pt(9)
    doc.add_paragraph()

def _render_paragraph_section(doc, title, content):
    doc.add_paragraph(title.upper(), style="Heading 1")
    if isinstance(content, list):
        for paragraph in content:
            p = doc.add_paragraph(str(paragraph))
            for run in p.runs: run.font.name, run.font.size = 'Arial', Pt(9)
    else:
        p = doc.add_paragraph(str(content))
        for run in p.runs: run.font.name, run.font.size = 'Arial', Pt(9)
    doc.add_paragraph()

def is_section_populated(data, section_name, section_conf=None, render_placeholders=False):
    if render_placeholders: return True
    if section_name == "Career History":
        if data.get("current_employment") or data.get("previous_employment") or data.get("career_history"):
            return True
        return False
    if section_conf:
        data_key = section_conf.get("data_key") or section_conf.get("key")
        if data_key:
            val = style_utils._get_nested_value(data, data_key)
            if val is None: return False
            if isinstance(val, (list, dict)) and len(val) == 0: return False
            if isinstance(val, str) and not val.strip(): return False
            return True
    key = DATA_KEY_MAP.get(section_name)
    if not key: return True
    val = data.get(key)
    if not val: return False
    if isinstance(val, (list, dict)) and len(val) == 0: return False
    return True

def _render_section_from_config(doc, section_conf, data):
    name = section_conf.get("name") or section_conf.get("title")
    if not name: return
    section_type = section_conf.get("type", "legacy")
    if section_type == "legacy":
        if name in SECTION_MAP: SECTION_MAP[name](doc, data)
        return
    if section_type == "key_values":
        _render_key_value_section(doc, name, section_conf.get("fields", []), data)
        return
    if section_type == "table":
        list_data = style_utils._get_nested_value(data, section_conf.get("data_key"), []) or []
        _render_table_section(doc, name, list_data, section_conf.get("columns", []), min_rows=section_conf.get("min_rows", 0))
        return
    if section_type == "list":
        items = style_utils._get_nested_value(data, section_conf.get("data_key"), section_conf.get("items", []))
        _render_list_section(doc, name, items)
        return
    if section_type == "paragraph":
        content = style_utils._get_nested_value(data, section_conf.get("data_key"), section_conf.get("text", ""))
        _render_paragraph_section(doc, name, content)
        return
    if section_type == "career_history":
        add_career_history(doc, data)
        return

def add_custom_section(doc, section_data):
    name = section_data.get("section_name", "Additional Information")
    content_type = section_data.get("content_type", "paragraph")
    content = section_data.get("content")
    
    if not content: return

    doc.add_paragraph(name.upper(), style="Heading 1")
    
    if content_type == "list" and isinstance(content, list):
        for item in content:
            p = doc.add_paragraph(str(item), style='List Bullet')
            for run in p.runs: run.font.name, run.font.size = 'Arial', Pt(9)
    elif content_type == "table" and isinstance(content, list):
        # dynamic table detection
        if content and isinstance(content[0], dict):
            cols = list(content[0].keys())
            table = doc.add_table(rows=len(content)+1, cols=len(cols))
            style_utils.apply_table_styles(table)
            for i, col in enumerate(cols):
                cell = table.rows[0].cells[i]
                cell.text = col.upper()
                style_utils.set_cell_shading(cell, 'BEBEBE')
                for run in cell.paragraphs[0].runs: run.font.name, run.font.size, run.font.bold = 'Arial', Pt(9), True
            for i, row_data in enumerate(content):
                row = table.rows[i+1]
                for j, col in enumerate(cols):
                    cell = row.cells[j]
                    cell.text = str(row_data.get(col, ""))
                    for run in cell.paragraphs[0].runs: run.font.name, run.font.size = 'Arial', Pt(9)
        else:
             p = doc.add_paragraph(str(content))
             for run in p.runs: run.font.name, run.font.size = 'Arial', Pt(9)
    else:
        # default to paragraph
        if isinstance(content, list): content = "\n".join(map(str, content))
        p = doc.add_paragraph(str(content))
        for run in p.runs: run.font.name, run.font.size = 'Arial', Pt(9)
    doc.add_paragraph()

# ------------------------------------------------
# GENERATE CV
# ------------------------------------------------

def generate_cv(data, output_path, section_order=None, layout=None, render_placeholders=False):
    doc = Document()
    style_utils.setup_sections(doc)
    style_utils.setup_body_text_style(doc)
    style_utils.setup_heading_style(doc)
    style_utils.setup_list_bullet_style(doc)

    generate_first_page(doc, data)
    doc.add_page_break()

    layout_name = layout or data.get("layout") or data.get("template") or data.get("layout_name")
    if not layout_name:
        layout_name, _ = style_utils.infer_layout_from_data(data)
    
    layout_conf = style_utils.load_layout_config(layout_name)
    section_items = []
    if layout_conf:
        if isinstance(layout_conf.get("sections"), list) and layout_conf.get("sections"):
            section_items = layout_conf.get("sections")
        else:
            section_order = section_order or layout_conf.get("section_order")
            enabled = layout_conf.get("enabled_sections", {})
            if section_order:
                section_items = [{"name": name, "enabled": enabled.get(name, True)} for name in section_order]

    if not section_items:
        section_order = section_order or [
            "Personal Details", "Courses", "Qualification", "Professional Membership", "Projects", 
            "Course Work", "Training", "Career Highlights", "Career Summary", "Career History"
        ]
        section_items = [{"name": n, "enabled": True} for n in section_order]

    for section in section_items:
        if isinstance(section, str): section = {"name": section, "enabled": True}
        name = section.get("name")
        if not name or not section.get("enabled", True): continue
        if not is_section_populated(data, name, section, render_placeholders=render_placeholders): continue
        _render_section_from_config(doc, section, data)

    # NEW: custom_sections from universal extraction
    custom_sections = data.get("custom_sections", [])
    if isinstance(custom_sections, list):
        for cs in custom_sections:
            add_custom_section(doc, cs)

    doc.save(output_path)
    print(f"Professional CV generated: {output_path}")

def create_template():
    output_file = "uploads/Professional_CV_Template.docx"
    generate_cv({}, output_file, render_placeholders=True)
    print(f"Professional Template created successfully: {output_file}")

def load_layout_keywords():
    return style_utils.load_layout_keywords()

def save_layout_keywords(keywords):
    return style_utils.save_layout_keywords(keywords)

def get_layout_preview(layout_name):
    conf = style_utils.load_layout_config(layout_name)
    if not conf: return None
    return conf.get("sections")

if __name__ == "__main__":
    create_template()