from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import json
import os

def _get_nested_value(data, key, default=None):
    if key is None: return default
    if isinstance(key, str) and "." in key:
        parts = key.split(".")
        curr = data
        for part in parts:
            if isinstance(curr, dict): curr = curr.get(part, None)
            else: return default
        return curr if curr is not None else default
    if isinstance(data, dict): return data.get(key, default)
    return default

def load_layout_config(layout_name_or_path):
    if not layout_name_or_path: return None
    path = layout_name_or_path
    if not os.path.isabs(path):
        # Look in the layouts directory relative to THIS file (style_utils.py)
        base_dir = os.path.dirname(os.path.abspath(__file__))
        path = os.path.join(base_dir, "layouts", f"{layout_name_or_path}.json")
        if not os.path.exists(path):
            path = os.path.join(base_dir, "layouts", layout_name_or_path)
    try:
        with open(path, "r", encoding="utf-8") as f: return json.load(f)
    except: return None

def infer_layout_from_data(data):
    position = _get_nested_value(data, "POSITION") or _get_nested_value(data, "header.position_applied")
    if not position: return None, None
    pos_lower = str(position).lower()
    base_dir = os.path.dirname(os.path.abspath(__file__))
    path = os.path.join(base_dir, "layouts", "layout_keywords.json")
    try:
        with open(path, "r", encoding="utf-8") as f: keywords = json.load(f) or {}
    except: return None, None
    
    import re
    for layout_name, kws in keywords.items():
        for kw in kws or []:
            if not kw: continue
            kw_lower = kw.lower()
            # If keyword is short (<= 3 chars), use word boundaries
            if len(kw_lower) <= 3:
                if re.search(rf'\b{re.escape(kw_lower)}\b', pos_lower):
                    return layout_name, kw
            else:
                if kw_lower in pos_lower:
                    return layout_name, kw
    return None, None

def get_sections_from_layout(data, layout_name=None):
    """
    Returns a list of section names based on the layout configuration.
    """
    if not layout_name:
        layout_name, _ = infer_layout_from_data(data)
    
    layout_conf = load_layout_config(layout_name)
    if not layout_conf:
        return None
        
    if isinstance(layout_conf.get("sections"), list) and layout_conf.get("sections"):
        return [s.get("name") if isinstance(s, dict) else s for s in layout_conf["sections"]]
    
    return None

def load_layout_keywords():
    """Loads the layout keywords configuration."""
    base_dir = os.path.dirname(os.path.abspath(__file__))
    path = os.path.join(base_dir, "layouts", "layout_keywords.json")
    try:
        with open(path, "r", encoding="utf-8") as f:
            return json.load(f) or {}
    except:
        return {}

def save_layout_keywords(keywords):
    """Saves the layout keywords configuration."""
    base_dir = os.path.dirname(os.path.abspath(__file__))
    path = os.path.join(base_dir, "layouts", "layout_keywords.json")
    try:
        with open(path, "w", encoding="utf-8") as f:
            json.dump(keywords, f, indent=2)
        return True
    except:
        return False

def set_cell_border(cell, top=True, bottom=True, left=True, right=True, sz='4', val='single', color='000000'):
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    
    existing_borders = tcPr.find(qn('w:tcBorders'))
    if existing_borders is not None:
        tcPr.remove(existing_borders)
        
    tcBorders = OxmlElement('w:tcBorders')
    
    borders = {'top': top, 'bottom': bottom, 'left': left, 'right': right}
    
    for border_name, apply in borders.items():
        border = OxmlElement(f'w:{border_name}')
        if apply:
            border.set(qn('w:val'), val)
            border.set(qn('w:sz'), sz)
            border.set(qn('w:color'), color)
        else:
            border.set(qn('w:val'), 'none')
        tcBorders.append(border)
            
    tcPr.append(tcBorders)

def set_cell_shading(cell, fill_color):
    if not fill_color:
        return
    
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    
    existing_shd = tcPr.find(qn('w:shd'))
    if existing_shd is not None:
        tcPr.remove(existing_shd)
        
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill_color)
    tcPr.append(shading_elm)

def format_run(run, font_name='Arial', font_size=9, bold=False, color=None, underline=False):
    run.font.name = font_name
    if font_size:
        run.font.size = Pt(font_size)
    run.font.bold = bold
    if underline:
        run.font.underline = True
    if color:
        run.font.color.rgb = color

def format_paragraph(paragraph, alignment=None, indent=None, spacing_before=None, spacing_after=None):
    if alignment is not None:
        paragraph.alignment = alignment
    if indent is not None:
        paragraph.paragraph_format.left_indent = Pt(indent)
    if spacing_before is not None:
        paragraph.paragraph_format.space_before = Pt(spacing_before)
    if spacing_after is not None:
        paragraph.paragraph_format.space_after = Pt(spacing_after)

def format_header_cell(cell, text, shade_color="BEBEBE", align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = text
    set_cell_shading(cell, shade_color)
    
    p = cell.paragraphs[0]
    format_paragraph(p, alignment=align)
    p.paragraph_format.left_indent = Pt(2)
    
    if p.runs:
        format_run(p.runs[0], bold=True)
    else:
        run = p.add_run(text)
        format_run(run, bold=True)
        
    set_cell_border(cell)

def format_data_cell(cell, text, shade_color=None, align=WD_ALIGN_PARAGRAPH.LEFT, bold=False):
    cell.text = str(text)
    if shade_color:
        set_cell_shading(cell, shade_color)
    
    p = cell.paragraphs[0]
    format_paragraph(p, alignment=align)
    
    if p.runs:
        format_run(p.runs[0], bold=bold)
    
    set_cell_border(cell)

def set_page_borders(doc):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    for section in doc.sections:
        sectPr = section._sectPr
        
        pgBorders = OxmlElement('w:pgBorders')
        pgBorders.set(qn('w:offsetFrom'), 'page')
        
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'double')
            border.set(qn('w:sz'), '24')                 
            border.set(qn('w:color'), '505050')                
            border.set(qn('w:space'), '24')                              
            pgBorders.append(border)
            
        sectPr.append(pgBorders)

def setup_sections(doc):
    from docx.shared import Inches
    for section in doc.sections:
        section.header_distance = Inches(0.64)
        section.footer_distance = Inches(0.81)
        
        section.left_margin = Inches(0.45)
        section.right_margin = Inches(0.45)
        section.top_margin = Inches(0.45)
        section.bottom_margin = Inches(0.45)
    set_page_borders(doc)

def apply_table_styles(table, row_height=0.15, table_width=6.87, vertical_align=WD_ALIGN_VERTICAL.TOP):
    num_cols = len(table.columns)
    table.width = Inches(table_width)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for row in table.rows:
        row.height = Inches(row_height)
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        for i, cell in enumerate(row.cells):
            if num_cols == 2:
                cell.width = Inches(2.12) if i == 0 else Inches(4.75)
            else:
                cell.width = Inches(table_width / num_cols)
            cell.vertical_alignment = vertical_align
            
            for p in cell.paragraphs:
                format_paragraph(p, indent=7.2, spacing_before=2, spacing_after=2)                

def setup_body_text_style(doc):
    style = doc.styles['Normal']
    
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    style.paragraph_format.left_indent = Pt(0)
    style.paragraph_format.right_indent = Pt(0)
    style.paragraph_format.first_line_indent = Pt(0)
    
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    
    style.paragraph_format.line_spacing = 1.15

def setup_heading_style(doc):
    style = doc.styles['Heading 1']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(9)
    font.underline = True
    font.bold = True
    font.color.rgb = RGBColor(0, 0, 0)
    
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    style.paragraph_format.space_before = Pt(4.7)
    style.paragraph_format.space_after = Pt(0.3)
    style.paragraph_format.left_indent = Inches(0.25)

def setup_list_bullet_style(doc):
    style = doc.styles['List Bullet']
    
    style.paragraph_format.left_indent = Inches(0.5)
    style.paragraph_format.first_line_indent = Inches(-0.25)
    
    style.paragraph_format.line_spacing = Pt(10.95)
    
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)

def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')
    rPr.append(color)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    b = OxmlElement('w:b')
    rPr.append(b)
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink

def add_common_footer(doc):
    section = doc.sections[0]
    footer = section.footer
    for p in footer.paragraphs:
        p.text = ""
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.left_indent = Inches(0.25)
    run = paragraph.add_run("SISOL RECRUITMENT")
    run.font.name = 'Arial'
    run.font.size = Pt(8)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    
    paragraph = footer.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.25)
    run = paragraph.add_run("Contact: 0102252001 for further details")
    run.font.name = 'Arial'
    run.font.size = Pt(8)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    
    paragraph = footer.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.25)
    run = paragraph.add_run("Email: ")
    run.font.name = 'Arial'
    run.font.size = Pt(8)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    add_hyperlink(paragraph, "hr@sisolprojects.co.za", "mailto:hr@sisolprojects.co.za")
    run = paragraph.add_run(" www.sisolrecruitment.co.za")
    run.font.name = 'Arial'
    run.font.size = Pt(8)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)

def add_page_header(doc):
    section = doc.sections[0]
    header = section.header
    for p in header.paragraphs:
        p.text = ""
    paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture('logo.png', width=Inches(1.3), height=Inches(0.78))
    
    paragraph = header.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("RECRUITMENT/ CV EXPRESS/ HUMAN RESOURCES CONSULTING")
    run.font.name = 'Arial'
    run.font.size = Pt(8)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    paragraph.paragraph_format.space_after = Pt(12)

def build_career_history_table(doc, title, history, style_variant="master"):
    if not history:
         history = [{"company": f"{{{{{title}_CO}}}}", "date": f"{{{{{title}_DT}}}}", "final_position": f"{{{{{title}_POS}}}}", "reason_for_leaving": "", "responsibilities": []}]
    
    table = doc.add_table(rows=1+(len(history)*5), cols=2)
    apply_table_styles(table)
    h_cell = table.rows[0].cells[0].merge(table.rows[0].cells[1])
    h_cell.text = title.upper()
    set_cell_shading(h_cell, 'BEBEBE')
    set_cell_border(h_cell)
    h_para = h_cell.paragraphs[0]
    h_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in h_para.runs: run.font.name, run.font.size, run.font.bold = 'Arial', Pt(9), True
    
    for i, job in enumerate(history):
        base = 1 + (i * 5)
        labels = ["Company", "Date", "Final Position", "Reason for Leaving"]
        vals = [job.get("company", ""), job.get("date", ""), job.get("final_position", ""), job.get("reason_for_leaving", "")]
        for idx, (lbl, val) in enumerate(zip(labels, vals)):
            row = table.rows[base + idx]
            c0, c1 = row.cells
            c0.text = lbl
            set_cell_shading(c0, 'F1F1F1')
            for run in c0.paragraphs[0].runs: run.font.name, run.font.size, run.font.bold = 'Arial', Pt(9), True
            c1.text = str(val).upper() if idx == 0 and style_variant == "master" else str(val)
            if style_variant in ("standard", "new"): set_cell_shading(c1, 'F1F1F1')
            for run in c1.paragraphs[0].runs: run.font.name, run.font.size, run.font.bold = 'Arial', Pt(9), (idx==0)
            set_cell_border(c0, top=(idx==0), bottom=(idx==3))
            set_cell_border(c1, top=(idx==0), bottom=(idx==3))
            
        b_cell = table.rows[base+4].cells[0].merge(table.rows[base+4].cells[1])
        resps = job.get("responsibilities", [])

        if isinstance(resps, list):
            import re as _re
            if style_variant == "new":
                _normalized = []
                for _r in resps:
                    if isinstance(_r, str):
                        _m = _re.match(r"^\s*(?P<title>[^():\n]{1,80}?)\s*\((?P<dates>[^)]+)\)\s*[:\-–—]\s*(?P<text>.+)$", _r)
                        if _m:
                            _title = _m.group('title').strip()
                            _dates = _m.group('dates').strip()
                            _text = _m.group('text').strip()
                            _bullets = [t.strip() for t in _re.split(r";|\.|\n", _text) if t.strip()]
                            _normalized.append({'title': _title, 'dates': _dates, 'bullets': _bullets})
                            continue
                        _m2 = _re.match(r"^\s*(?P<header>[^:]{1,60}?)\s*:\s*(?P<body>.+)$", _r)
                        if _m2 and _re.search(r"\d{4}", _m2.group('body')):
                            _header = _m2.group('header').strip()
                            _body = _m2.group('body').strip()
                            _date_match = _re.search(r"(\b\d{4}[^,;\n]*)", _body)
                            _dates = _date_match.group(1).strip() if _date_match else ""
                            _remainder = _re.sub(_re.escape(_dates), '', _body).strip() if _dates else _body
                            _bullets = [t.strip() for t in _re.split(r";|\.|\n", _remainder) if t.strip()]
                            _normalized.append({'title': _header, 'dates': _dates, 'bullets': _bullets})
                            continue
                    _normalized.append(_r)
                resps = _normalized

            for r in resps:
                if isinstance(r, dict):
                    hp = b_cell.add_paragraph()
                    if style_variant == "new": hp.paragraph_format.space_before = Pt(6)
                    hdr_text = f"{r.get('title','')}"
                    if r.get('dates'): hdr_text += f" - {r.get('dates','')}"
                    hp.add_run(hdr_text).font.name, hp.runs[0].font.size, hp.runs[0].font.bold = 'Arial', Pt(9), True
                    for b in r.get("bullets", []):
                        bp = b_cell.add_paragraph(str(b), style='List Bullet')
                        for run in bp.runs: run.font.name, run.font.size = 'Arial', Pt(9)

                    sub_ach = r.get("achievements", [])
                    if sub_ach:
                        ach_hdr = b_cell.add_paragraph()
                        ach_hdr.add_run("Achievements:").font.name, ach_hdr.runs[0].font.size, ach_hdr.runs[0].font.bold = 'Arial', Pt(9), True
                        for a in sub_ach:
                            ap = b_cell.add_paragraph(str(a), style='List Bullet')
                            for run in ap.runs: run.font.name, run.font.size = 'Arial', Pt(9)
                else:
                    bp = b_cell.add_paragraph(str(r), style='List Bullet')
                    for run in bp.runs: run.font.name, run.font.size = 'Arial', Pt(9)
        else:
            p = b_cell.add_paragraph(str(resps))
            for run in p.runs: run.font.name, run.font.size = 'Arial', Pt(9)

        job_ach = job.get("achievements", [])
        if job_ach:
            ach_hdr = b_cell.add_paragraph()
            ach_hdr.add_run("Key Achievements:").font.name, ach_hdr.runs[0].font.size, ach_hdr.runs[0].font.bold = 'Arial', Pt(9), True
            for a in job_ach:
                ap = b_cell.add_paragraph(str(a), style='List Bullet')
                for run in ap.runs: run.font.name, run.font.size = 'Arial', Pt(9)

        set_cell_border(b_cell)
    doc.add_paragraph()
