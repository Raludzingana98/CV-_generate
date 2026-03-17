from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import style_utils
import datetime
import json

# ------------------------------------------------
# FIRST PAGE (COVER PAGE)
# ------------------------------------------------
def generate_first_page(doc, data=None):
    data = data or {}
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("EMAIL: ")
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    run.font.bold = True
    style_utils.add_hyperlink(paragraph, "HR@SISOLPROJECTS.CO.ZA", "mailto:INFO@SISOL-PROJECTS.CO.ZA")
    for run in paragraph.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(10)
        run.font.bold = True
    paragraph.paragraph_format.space_after = Pt(189)
    doc.add_paragraph()
    
    table = doc.add_table(rows=2, cols=1)
    style_utils.apply_table_styles(table)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.width = Inches(6.77)
    
    name_row = table.rows[0]
    name_cell = name_row.cells[0]
    name_para = name_cell.paragraphs[0]
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_val = data.get("CANDIDATE_NAME", "{{CANDIDATE_NAME}}")
    run = name_para.add_run(f"CV OF: {str(name_val).upper()}")
    run.font.name = 'Arial'
    run.font.size = Pt(28)
    run.font.bold = True
    style_utils.set_cell_shading(name_cell, 'D9D9D9')
    
    pos_row = table.rows[1]
    pos_cell = pos_row.cells[0]
    pos_para = pos_cell.paragraphs[0]
    pos_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pos_val = data.get("POSITION", "{{POSITION}}")
    run = pos_para.add_run(f"Position Applied for: {pos_val}")
    run.font.name = 'Arial'
    run.font.size = Pt(22)
    run.font.bold = True
    style_utils.set_cell_shading(pos_cell, 'F1F1F1')
    
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rec_val = data.get("RECRUITER_NAME", "{{RECRUITER_NAME}}")
    run = paragraph.add_run(f"CV Presented by {rec_val}")
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(23, 54, 93)
    
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    current_date = f"{datetime.datetime.now().day} {datetime.datetime.now().strftime('%B %Y')}"
    date_val = data.get("DATE", current_date)
    run = paragraph.add_run(date_val)
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(23, 54, 93)
    paragraph.paragraph_format.space_after = Pt(30)
    
    table = doc.add_table(rows=1, cols=1)
    style_utils.apply_table_styles(table)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.width = Inches(3.3)
    table.rows[0].height = Inches(1.7)

    cell = table.rows[0].cells[0]
    cell.width = Inches(3.3)
    style_utils.set_cell_border(cell)
    style_utils.set_cell_shading(cell, 'D9D9D9')

    def _add_centered(txt, bold=True):
        p = cell.add_paragraph(txt)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.runs[0]
        r.font.name = 'Arial'
        r.font.size = Pt(9)
        r.font.bold = bold
        return p

    _add_centered("Previous Salary")
    _add_centered(data.get("PREVIOUS_SALARY", "{{PREVIOUS_SALARY}}"), bold=False)

    sp = cell.add_paragraph()
    sp.paragraph_format.space_after = Pt(2)

    _add_centered("Salary Expectation")
    _add_centered(data.get("SALARY_EXPECTATION", "{{SALARY_EXPECTATION}}"), bold=False)
    doc.add_paragraph()

# ------------------------------------------------
# SECTION RENDERING HELPERS
# ------------------------------------------------

def add_personal_details(doc, data):
    style_utils.add_page_header(doc)
    table = doc.add_table(rows=16, cols=2)
    style_utils.apply_table_styles(table)
    header_row = table.rows[0]
    header_cell = header_row.cells[0].merge(header_row.cells[1])
    header_para = header_cell.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header_para.add_run("PERSONAL DETAILS")
    run.font.name = 'Arial'
    run.font.size = Pt(9)
    run.font.bold = True
    style_utils.set_cell_shading(header_cell, 'BEBEBE')
    
    fields = [
        ("Full Name", "CANDIDATE_NAME"),
        ("Known As", "KNOWN_AS"),
        ("Race and Gender", "RACE_GENDER"),
        ("Nationality", "NATIONALITY"),
        ("Marital Status", "MARITAL_STATUS"),
        ("Date of birth", "DOB"),
        ("Driver's License / Own Car", "DRIVERS"),
        ("Notice Period", "NOTICE"),
        ("Do you have clear criminal record", "CRIMINAL"),
        ("Area of residence", "RESIDENCE"),
        ("Languages", "LANGUAGES"),
        ("Salary Expectations", "SALARY_EXP"),
        ("Skills", "SKILLS"),
        ("Computer literacy", "COMPUTER"),
        ("Experience", "EXPERIENCE_SUMMARY")
    ]
    for idx, (label, key) in enumerate(fields):
        row = table.rows[idx + 1]
        c0, c1 = row.cells
        c0.text = label
        style_utils.set_cell_shading(c0, 'F1F1F1')
        p0 = c0.paragraphs[0]
        for run in p0.runs:
            run.font.name, run.font.size, run.font.bold = 'Arial', Pt(9), True
        
        c1.text = str(data.get(key, f"{{{{{key}}}}}"))
        p1 = c1.paragraphs[0]
        for run in p1.runs:
            run.font.name, run.font.size, run.font.bold = 'Arial', Pt(9), False
    
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            style_utils.set_cell_border(cell, top=(r_idx==0), bottom=True, left=(c_idx==0), right=(c_idx==len(row.cells)-1 or (r_idx==0 and c_idx==0)))
    doc.add_paragraph()

def add_qualifications(doc, data):
    doc.add_paragraph("QUALIFICATIONS", style="Heading 1")
    quals = data.get("qualifications", [])
    num_rows = max(len(quals), 4) + 1
    table = doc.add_table(rows=num_rows, cols=4)
    style_utils.apply_table_styles(table)
    headers = ["QUALIFICATION", "INSTITUTION", "YEAR", "STATUS"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        style_utils.set_cell_shading(cell, 'BEBEBE')
        style_utils.set_cell_border(cell)
        for run in cell.paragraphs[0].runs: run.font.name, run.font.size, run.font.bold = 'Arial', Pt(9), True
            
    for i in range(1, num_rows):
        row = table.rows[i]
        if i-1 < len(quals):
            q = quals[i-1]
            vals = [q.get("qualification", ""), q.get("institution", ""), q.get("date", ""), q.get("status", "")]
        else:
            vals = [f"{{{{QUAL_{i}}}}}", f"{{{{INSTITUTION_{i}}}}}", f"{{{{YEAR_{i}}}}}", f"{{{{STATUS_{i}}}}}" ]
        for j, val in enumerate(vals):
            cell = row.cells[j]
            cell.text = str(val)
            style_utils.set_cell_border(cell)
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT if j < 2 else WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs: run.font.name, run.font.size = 'Arial', Pt(9)
        style_utils.set_cell_shading(row.cells[0], 'F1F1F1')
    doc.add_paragraph()

def add_memberships(doc, data):
    doc.add_paragraph("PROFESSIONAL MEMBERSHIPS", style="Heading 1")
    members = data.get("professional_membership", [])
    num_rows = max(len(members), 2) + 1
    table = doc.add_table(rows=num_rows, cols=4)
    style_utils.apply_table_styles(table)
    headers = ["NAME", "PROFESSION", "YEAR", "PROF NO"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        style_utils.set_cell_shading(cell, 'BEBEBE')
        style_utils.set_cell_border(cell)
        for run in cell.paragraphs[0].runs: run.font.name, run.font.size, run.font.bold = 'Arial', Pt(9), True
            
    for i in range(1, num_rows):
        row = table.rows[i]
        if i-1 < len(members):
            m = members[i-1]
            vals = [m.get("board_name", ""), m.get("profession", ""), m.get("year", ""), m.get("prof_no", "")]
        else:
            vals = [f"{{{{MEMBER_NAME_{i}}}}}", f"{{{{PROFESSION_{i}}}}}", f"{{{{MEMBER_YEAR_{i}}}}}", f"{{{{PROF_NO_{i}}}}}" ]
        for j, val in enumerate(vals):
            cell = row.cells[j]
            cell.text = str(val)
            style_utils.set_cell_border(cell)
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT if j < 2 else WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs: run.font.name, run.font.size = 'Arial', Pt(9)
        style_utils.set_cell_shading(row.cells[0], 'F1F1F1')
    doc.add_paragraph()

def add_projects(doc, data):
    doc.add_paragraph("PROJECTS", style="Heading 1")
    projects = data.get("projects", [])
    if not projects:
        projects = [{"title": f"{{{{PROJECT_TITLE_{i}}}}}", "value": f"{{{{PROJECT_VALUE_{i}}}}}", "bullets": [f"{{{{PROJECT_DESCRIPTION_{i}}}}}"]} for i in range(1, 3)]
    for proj in projects:
        title = proj.get("title", proj.get("project_name", ""))
        value = proj.get("value", proj.get("project_value", ""))
        bullets = proj.get("bullets", proj.get("description", []))
        if isinstance(bullets, str): bullets = [bullets]
        
        p_title = doc.add_paragraph()
        p_title.add_run(str(title)).font.name, p_title.runs[0].font.size, p_title.runs[0].font.bold = 'Arial', Pt(9), True
        p_title.paragraph_format.left_indent = Inches(0.25)
        p_title.paragraph_format.space_after = Pt(0)
        
        p_value = doc.add_paragraph()
        p_value.add_run(str(value)).font.name, p_value.runs[0].font.size, p_value.runs[0].font.bold = 'Arial', Pt(9), True
        p_value.paragraph_format.left_indent = Inches(0.25)
        p_value.paragraph_format.space_after = Pt(0)
        
        for b in bullets:
            pb = doc.add_paragraph(str(b), style='List Bullet')
            for run in pb.runs: run.font.name, run.font.size = 'Arial', Pt(9)
        doc.add_paragraph()

def add_career_summary(doc, data):
    doc.add_paragraph("CAREER SUMMARY", style="Heading 1")
    summary = data.get("career_summary", [])
    num_rows = max(len(summary), 5) + 1
    table = doc.add_table(rows=num_rows, cols=3)
    style_utils.apply_table_styles(table)
    headers = ["COMPANY", "POSITION", "DATES"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        style_utils.set_cell_shading(cell, 'BEBEBE')
        style_utils.set_cell_border(cell)
        for run in cell.paragraphs[0].runs: run.font.name, run.font.size, run.font.bold = 'Arial', Pt(9), True
            
    for i in range(1, num_rows):
        row = table.rows[i]
        if i-1 < len(summary):
            s = summary[i-1]
            vals = [s.get("company", ""), s.get("position", ""), s.get("dates", "")]
        else:
            vals = [f"{{{{SUMMARY_COMPANY_{i}}}}}", f"{{{{SUMMARY_POSITION_{i}}}}}", f"{{{{SUMMARY_DATES_{i}}}}}" ]
        for j, val in enumerate(vals):
            cell = row.cells[j]
            cell.text = str(val)
            style_utils.set_cell_border(cell)
            for run in cell.paragraphs[0].runs: run.font.name, run.font.size = 'Arial', Pt(9)
    doc.add_paragraph()

def add_career_history(doc, data):
    doc.add_paragraph("CAREER HISTORY", style="Heading 1")
    style_utils.build_career_history_table(doc, "CURRENT EMPLOYMENT", data.get("current_employment", []), style_variant="master")
    style_utils.build_career_history_table(doc, "PREVIOUS EMPLOYMENT", data.get("previous_employment", []), style_variant="master")

def add_custom_section(doc, section_data):
    name = section_data.get("section_name", "Additional Information")
    content_type = section_data.get("content_type", "paragraph")
    content = section_data.get("content")
    
    if not content: return

    doc.add_paragraph(name.upper(), style="Heading 1")
    
    if content_type == "list" and isinstance(content, list):
        for item in content:
            p = doc.add_paragraph(str(item), style='List Bullet')
            for run in p.runs: run.font.name, run.font.size = 'Arial', Pt(9)
    elif content_type == "table" and isinstance(content, list):
        if content and isinstance(content[0], dict):
            cols = list(content[0].keys())
            table = doc.add_table(rows=len(content)+1, cols=len(cols))
            style_utils.apply_table_styles(table)
            for i, col in enumerate(cols):
                cell = table.rows[0].cells[i]
                cell.text = col.upper()
                style_utils.set_cell_shading(cell, 'BEBEBE')
                for run in cell.paragraphs[0].runs: run.font.name, run.font.size, run.font.bold = 'Arial', Pt(9), True
            for i, row_data in enumerate(content):
                row = table.rows[i+1]
                for j, col in enumerate(cols):
                    cell = row.cells[j]
                    cell.text = str(row_data.get(col, ""))
                    for run in cell.paragraphs[0].runs: run.font.name, run.font.size = 'Arial', Pt(9)
        else:
             p = doc.add_paragraph(str(content))
             for run in p.runs: run.font.name, run.font.size = 'Arial', Pt(9)
    else:
        if isinstance(content, list): content = "\n".join(map(str, content))
        p = doc.add_paragraph(str(content))
        for run in p.runs: run.font.name, run.font.size = 'Arial', Pt(9)
    doc.add_paragraph()

def add_highlights(doc, data):
    doc.add_paragraph("CAREER HIGHLIGHTS", style="Heading 1")
    highlights = data.get("career_highlights", [])
    if not highlights: highlights = ["{{HIGHLIGHT_1}}", "{{HIGHLIGHT_2}}", "{{HIGHLIGHT_3}}"]
    for h in highlights:
        p = doc.add_paragraph(str(h), style='List Bullet')
        for run in p.runs: run.font.name, run.font.size = 'Arial', Pt(9)
    doc.add_paragraph()

def add_course_work(doc, data):
    doc.add_paragraph("COURSE WORK", style="Heading 1")
    course_work = data.get("course_work", [])
    if not course_work: course_work = ["{{COURSE_WORK_1}}", "{{COURSE_WORK_2}}"]
    for cw in course_work:
        p = doc.add_paragraph(str(cw), style='List Bullet')
        for run in p.runs: run.font.name, run.font.size = 'Arial', Pt(9)
    doc.add_paragraph()

def add_training(doc, data):
    doc.add_paragraph("TRAINING", style="Heading 1")
    training = data.get("training", [])
    if not training: training = ["{{TRAINING_1}}", "{{TRAINING_2}}"]
    for t in training:
        p = doc.add_paragraph(str(t), style='List Bullet')
        for run in p.runs: run.font.name, run.font.size = 'Arial', Pt(9)
    doc.add_paragraph()

def add_courses(doc, data):
    doc.add_paragraph("COURSES", style="Heading 1")
    courses = data.get("courses", [])
    num_rows = max(len(courses), 4) + 1
    table = doc.add_table(rows=num_rows, cols=4)
    style_utils.apply_table_styles(table)
    headers = ["COURSE NAME", "INSTITUTION", "DATE", "STATUS"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        style_utils.set_cell_shading(cell, 'BEBEBE')
        style_utils.set_cell_border(cell)
        for run in cell.paragraphs[0].runs: run.font.name, run.font.size, run.font.bold = 'Arial', Pt(9), True
            
    for i in range(1, num_rows):
        row = table.rows[i]
        if i-1 < len(courses):
            c = courses[i-1]
            vals = [c.get("course_name", ""), c.get("institution", ""), c.get("date", ""), c.get("status", "")]
        else:
            vals = [f"{{{{COURSE_{i}}}}}", f"{{{{INSTITUTION_{i}}}}}", f"{{{{DATE_{i}}}}}", f"{{{{STATUS_{i}}}}}" ]
        for j, val in enumerate(vals):
            cell = row.cells[j]
            cell.text = str(val)
            style_utils.set_cell_border(cell)
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT if j < 2 else WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs: run.font.name, run.font.size = 'Arial', Pt(9)
        style_utils.set_cell_shading(row.cells[0], 'F1F1F1')
    doc.add_paragraph()

# ------------------------------------------------
# CONFIGURATION
# ------------------------------------------------

SECTION_MAP = {
    "Personal Details": add_personal_details,
    "Qualification": add_qualifications,
    "Professional Membership": add_memberships,
    "Projects": add_projects,
    "Career Summary": add_career_summary,
    "Career History": add_career_history,
    "Career Highlights": add_highlights,
    "Course Work": add_course_work,
    "Training": add_training,
    "Courses": add_courses
}

DATA_KEY_MAP = {
    "Personal Details": "personal_details",
    "Qualification": "qualifications",
    "Professional Membership": "professional_membership",
    "Projects": "projects",
    "Career Summary": "career_summary",
    "Career History": "career_history",
    "Career Highlights": "career_highlights",
    "Course Work": "course_work",
    "Training": "training",
    "Courses": "courses"
}

def is_section_populated(data, section_name):
    if section_name == "Career History":
        if data.get("current_employment") or data.get("previous_employment") or data.get("career_history"):
            return True
        return False
    key = DATA_KEY_MAP.get(section_name)
    if not key: return True
    val = data.get(key)
    if not val: return False
    if isinstance(val, (list, dict)) and len(val) == 0: return False
    return True

def generate_cv(data, output_path, section_order=None, layout=None):
    doc = Document()
    style_utils.setup_sections(doc)
    style_utils.setup_body_text_style(doc)
    style_utils.setup_heading_style(doc)
    style_utils.setup_list_bullet_style(doc)
    
    generate_first_page(doc, data)
    doc.add_page_break()
    
    if not section_order:
        section_order = style_utils.get_sections_from_layout(data, layout)
        
    if not section_order:
        section_order = [
            "Personal Details", "Courses", "Qualification", "Professional Membership", "Projects", 
            "Course Work", "Training", "Career Highlights", "Career Summary", "Career History"
        ]
    
    for section_name in section_order:
        if section_name in SECTION_MAP:
            if is_section_populated(data, section_name):
                SECTION_MAP[section_name](doc, data)
            else:
                print(f"Skipping empty section: {section_name}")
        else:
            print(f"Warning: Section '{section_name}' not found in mapping.")
    
    # NEW: custom_sections from universal extraction
    custom_sections = data.get("custom_sections", [])
    if isinstance(custom_sections, list):
        for cs in custom_sections:
            add_custom_section(doc, cs)
    
    style_utils.add_common_footer(doc)
    doc.save(output_path)
    print(f"Master CV generated: {output_path}")

def create_template():
    output_file = "uploads/Master_CV_Template.docx"
    generate_cv({}, output_file)
    print(f"Master Template created successfully: {output_file}")

if __name__ == "__main__":
    create_template()
